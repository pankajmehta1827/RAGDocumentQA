"""Extracts raw text from uploaded PDF or DOCX files."""
from __future__ import annotations

import io
import sys
from functools import lru_cache

import docx
import pymupdf
from pypdf import PdfReader

# EasyOCR's first-run model download prints a Unicode progress bar; on
# Windows, stdout often defaults to a legacy codepage (e.g. cp1252) that
# can't encode it and raises UnicodeEncodeError mid-download.
if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Return the plain text content of a PDF or DOCX file."""
    ext = filename.lower().rsplit(".", 1)[-1]

    if ext == "pdf":
        return _extract_pdf(file_bytes)
    if ext == "docx":
        return _extract_docx(file_bytes)

    raise ValueError(f"Unsupported file type: .{ext}. Only PDF and DOCX are supported.")


def _extract_pdf(file_bytes: bytes) -> str:
    """Try progressively more expensive extraction strategies until one
    yields text: pypdf, then PyMuPDF (a different parser that often
    succeeds on multi-column/custom-font resume templates pypdf can't
    read), then OCR as a last resort for genuinely image-only PDFs."""
    text = _extract_pdf_pypdf(file_bytes)
    if not text.strip():
        text = _extract_pdf_pymupdf(file_bytes)
    if not text.strip():
        text = _extract_pdf_ocr(file_bytes)
    return text


def _extract_pdf_pypdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages)


def _extract_pdf_pymupdf(file_bytes: bytes) -> str:
    with pymupdf.open(stream=file_bytes, filetype="pdf") as doc:
        return "\n\n".join(page.get_text() for page in doc)


@lru_cache(maxsize=1)
def _ocr_reader():
    import easyocr

    return easyocr.Reader(["en"], gpu=False, verbose=False)


def _extract_pdf_ocr(file_bytes: bytes) -> str:
    """Render each page to an image and run OCR. Used only when the PDF
    has no extractable text layer at all (e.g. a scanned document)."""
    try:
        reader = _ocr_reader()
    except Exception:
        return ""

    try:
        pages_text = []
        with pymupdf.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                pixmap = page.get_pixmap(dpi=150)
                lines = reader.readtext(pixmap.tobytes("png"), detail=0)
                pages_text.append("\n".join(lines))
                del pixmap
        return "\n\n".join(pages_text)
    except Exception:
        return ""


def _extract_docx(file_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)
